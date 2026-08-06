# -*- coding: utf-8 -*-
"""Build a Google Docs compatible .docx of the hackathon ideation doc."""
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x1A, 0x53, 0x6B)
MUTED = RGBColor(0x5F, 0x6B, 0x73)
LINK = RGBColor(0x0B, 0x57, 0xD0)

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
pf = normal.paragraph_format
pf.space_after = Pt(8)
pf.line_spacing = 1.15

for name, size, color, before, after in [
    ("Title", 26, ACCENT, 0, 4),
    ("Heading 1", 18, ACCENT, 20, 8),
    ("Heading 2", 14, ACCENT, 14, 6),
    ("Heading 3", 11.5, RGBColor(0x33, 0x33, 0x33), 10, 4),
]:
    st = doc.styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for s in doc.sections:
    s.left_margin = s.right_margin = Inches(0.9)
    s.top_margin = s.bottom_margin = Inches(0.8)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0B57D0")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(c)
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hl.append(new_run)
    paragraph._p.append(hl)
    return hl


def para(text="", style=None, bold=False, italic=False, size=None,
         color=None, space_after=None, align=None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        if size:
            r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p


def labelled(label, text):
    """Bold inline label followed by body text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(label + "  ")
    r.bold = True
    p.add_run(text)
    return p


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        shade(hdr[i], "E8EEF2")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(val))
            r.font.size = Pt(10)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def source_line(label, url):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.15)
    r = p.add_run(label + " — ")
    r.font.size = Pt(10)
    add_hyperlink(p, url, url)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


# ============================== TITLE ==============================
para("AI Hackathon Festival 2026", style="Title", space_after=2)
para("Ideation — University of Waikato", size=13, color=MUTED, space_after=2)
para("Two-day build · Teams of 3–7 · Six candidate projects",
     size=10.5, color=MUTED, italic=True, space_after=14)

# ============================== BRIEF ==============================
doc.add_heading("The brief", level=1)

doc.add_heading("Judging", level=2)
para("Four parts, weighted equally.", space_after=4)
make_table(
    ["Criterion", "What it covers"],
    [
        ["Inspiration", "A clear problem with real impact"],
        ["Technology", "Feasibility, quality, how you use AI"],
        ["Design", "Usability and originality"],
        ["Presentation", "Clear, structured, within 5 minutes"],
    ],
    widths=[1.5, 5.0],
)

doc.add_heading("Prizes", level=2)
make_table(
    ["Prize", "Value", "Basis"],
    [
        ["Local", "$250", "Best Waikato solution"],
        ["TAIAO", "$1,000", "SDG alignment, feasibility, impact"],
        ["Technical Brilliance", "$1,000", "Technical thinking and prototype execution"],
    ],
    widths=[1.7, 0.9, 3.9],
)
para("Local winners go through to national judging.", size=10, color=MUTED)

doc.add_heading("What scores well", level=2)
para("Judges reward execution over novelty. Working software aimed at a specific "
     "group beats a clever idea that only half runs. A plain LLM wrapper no longer "
     "stands out. Health and education chatbots are the most common entries, so they "
     "are the hardest lane to win in.")
para("One thing separates entries reliably: producing a number that maps to an "
     "official UN indicator. It gives judges something concrete to check.")

doc.add_heading("Context", level=2)
para("Only 16% of SDG targets are projected to be met by 2030. The UN's own count puts "
     "36% of 139 targets on track or progressing moderately. Four years left.")
para("Two of the three problems suggested in the event PDF — Facilities Helpdesk and "
     "Peer Support — are chatbot-shaped. Choosing one means competing mostly on "
     "presentation.")

doc.add_page_break()

# ============================== IDEA 1 ==============================
doc.add_heading("Idea 1 — Cold-chain monitor for last-mile clinics", level=1)

labelled("Problem",
         "Vaccines and insulin spoil quietly when a fridge fails or a transport box "
         "warms above 8 °C. Rural clinics often find out only when the stock is "
         "already unusable.")
labelled("SDG",
         "SDG 3 (Good Health). Protects the potency of health supplies reaching "
         "underserved patients. Indicator 3.b.1, vaccine coverage.")

doc.add_heading("Worth knowing first", level=2)
para("This device already exists commercially. Nexleaf ColdTrace X is WHO-prequalified "
     "and does everything an ESP32 rig would: ±0.5 °C accuracy, BLE sensors, cellular "
     "or WiFi sync, five-year warranty, and a backend running 15,000+ devices. "
     "Berlinger Fridge-tag is the standalone equivalent.")
para("But research on why cold chains fail points somewhere else. In one field study, "
     "91.7% of facilities had no cold chain policy and none had an equipment "
     "contingency plan. The alert fires and nobody knows who is meant to act. "
     "That gap is open.")

doc.add_heading("Approach", level=2)
para("Keep the ESP32 and temperature sensor, but treat it as the input, not the "
     "product. The product is what happens after an excursion is detected.", space_after=4)
numbered("Dose disposition. The system reads the excursion — how long, how hot or "
         "cold, which fridge — and checks it against vaccine stability rules to "
         "recommend use, quarantine, or discard for each product, showing its "
         "reasoning. Today this decision is made manually.")
numbered("Escalation. Identifies who is on call, whether the immunisation coordinator "
         "has been contacted, and drafts the incident record.")
numbered("Tamper-evident log. Events are hash-stamped so a flagged batch cannot be "
         "quietly un-flagged.")

doc.add_heading("New Zealand angle", level=2)
bullet("The local process is CCA (Cold Chain Accreditation), run by the Ministry of "
       "Health and valid up to three years. Pharmacies are audited separately by "
       "Medicines Control.")
bullet("The 2017 National Standards are under review by Health NZ right now.")
bullet("Current interim rule: fridges over ten years old are allowed only with "
       "continuous real-time monitoring and out-of-hours alarming. That is existing "
       "demand for this, today.")
bullet("Freezing, not overheating, causes most vaccine damage in NZ.")
bullet("New Zealand's cold chain programme cut vaccine wastage from 17% to 2% between "
       "2002 and 2008, saving an estimated $4M a year.")

doc.add_heading("Build and demo", level=2)
labelled("Offline-first", "Buffers locally, syncs on reconnect.")
labelled("Why 48 hours works",
         "The firmware is a read-sensor, append-log loop. The web layer is standard "
         "CRUD with a time-series chart. The decision layer is a prompt over a rules "
         "table, not a model you have to train.")
labelled("Demo",
         "Put the probe in a cup of ice and pull it out mid-pitch. The dashboard flags "
         "the excursion, then the system produces a disposition recommendation and a "
         "drafted incident report on screen.")

doc.add_heading("Risks", level=2)
bullet("Venue WiFi is unreliable. Use BLE-to-phone or a local access point as the main "
       "sync path.")
bullet("Frame output as decision support pending coordinator sign-off, not as a "
       "clinical ruling.")
bullet("Record a video of the excursion firing as a fallback.")

para("Solid if you pitch the decision layer. Weak if you pitch the logger.",
     italic=True, color=MUTED)

doc.add_page_break()

# ============================== IDEA 2 ==============================
doc.add_heading("Idea 2 — Surplus matching for growers", level=1)

labelled("Problem",
         "Produce is culled before packing for cosmetic reasons and sold by the bin, "
         "not the bag. It never enters a retail inventory system, so no existing "
         "food-waste platform can see it. Meanwhile household food waste in NZ is "
         "already falling — 12.2% to 10.9% of food bought, $3.2B to $3.0B — so the "
         "consumer end of the problem is shrinking while this end is not.")
labelled("SDG", "SDG 12.3 (halve food waste) and SDG 2.1 (end hunger). "
                "Two indicators, one build.")

doc.add_heading("The numbers", level=2)
bullet("NZ Food Network and Lineage redirected over 1,000 pallets — about 635,000 kg — "
       "of frozen vegetables rejected only for cosmetic irregularities.")
bullet("One in three NZ households faced food insecurity in the past year "
       "(Hunger Monitor 2025, published March 2026).")
bullet("The burden is uneven: Pacific peoples 64%, Māori 51%, ages 18–24 50%.")
bullet("Only 44% of food-insecure households have ever accessed food relief.")

doc.add_heading("Why this is open", level=2)
para("NZ food rescue matching is route-based and manual. NZFN's stated approach for "
     "bulk surplus is to act as a single point of contact — a human broker. Regular "
     "donors sit on fixed recurring routes.")
para("The existing apps solve adjacent problems: Olio is neighbour-to-neighbour, "
     "Too Good To Go and Flashfood work off retail shelves, Spoiler Alert handles "
     "packaged goods and inventory. Imperfect Foods and Misfits Market turned cosmetic "
     "rejects into a subscription box rather than a matching layer. None of them can "
     "see stock that is culled before it becomes inventory.")

doc.add_heading("Approach", level=2)
numbered("A grower photographs a bin. A vision model grades the reject class and "
         "estimates volume. Manual data entry is what stops growers participating, so "
         "removing it matters more than it sounds.")
numbered("A matcher scores allocations against shelf life, cold storage availability, "
         "hub capacity, travel time, and need severity. This is a scoring function, "
         "not a chatbot.")
numbered("An impact ledger tracks kg diverted, meal equivalents, and CO₂e avoided, "
         "mapped to SDG 12.3.")

doc.add_heading("Domain detail and scale", level=2)
para("The Food Act 2014 excludes food with compromised sanitary seals unless it is "
     "heated. Naming a real constraint like this shows you looked at the sector.")
para("64 food hubs, 2,000+ frontline groups, and roughly 28 million meal equivalents a "
     "year are already moving. You would be adding a layer to an existing network "
     "rather than building one.")

doc.add_heading("Build and demo", level=2)
labelled("Why 48 hours works",
         "Vision grading is one API call with a rubric. A greedy allocator demos the "
         "same as an optimal one at this scale. The dashboard is CRUD, a map, and a "
         "chart.")
labelled("Demo",
         "Photograph a crate of misshapen carrots. It gets graded, the matcher "
         "allocates across three mock hubs, and the ledger updates. Then mark a hub's "
         "chiller offline and let it reallocate.")

doc.add_heading("Risks", level=2)
bullet("A two-sided marketplace with thin supply fails. Pitch it as a tool for one "
       "broker who already has both sides, not as a marketplace.")
bullet("Position it as complementary to NZFN rather than competing.")

para("The best-balanced option: local data, a real gap, AI that is not a chatbot, two "
     "indicators, and a mechanism that works beyond NZ.", italic=True, color=MUTED)

doc.add_page_break()

# ============================== IDEA 3 ==============================
doc.add_heading("Idea 3 — Disability funding reassessment tool", level=1)

labelled("Problem",
         "NZ's disability support funding is changing now, and the people affected "
         "cannot work out what it means for them. From 1 April 2026 existing NASC users "
         "get a set budget with a limit. From 1 October 2026 reassessments run under "
         "the new system. Budgets are based on past spending, so anyone who underspent "
         "during a restrictive period risks being held to that lower figure.")
labelled("SDG", "SDG 10.2 (social and economic inclusion regardless of disability).")

doc.add_heading("Why the timing matters", level=2)
bullet("Waikato is an EGL site. EGL Waikato and Mana Whaikaha whānau keep direct "
       "funding. The event is inside the affected region and the judges are local.")
bullet("Reassessments start two months after the hackathon.")
bullet("The policy is actively contested. Blind Low Vision NZ has flagged clauses 7 "
       "and 8 of the Disability Support Services Bill 2026 as a step away from EGL "
       "principles, framing support as \u201cwithin available funding\u201d and putting "
       "family and community ahead of State support.")
bullet("CCS Disability Action says Budget 2026 squeezes community-based support while "
       "protecting residential care.")

doc.add_heading("Approach", level=2)
numbered("Budget simulator. Enter current supports, model the allocation under the new "
         "past-spending rule, and show the difference — flagging where the old March "
         "2024 purchasing restrictions may have artificially lowered someone's "
         "baseline. That is the specific harm advocates are warning about.")
numbered("Evidence pack generator. Retrieval over the National Standards, the February "
         "2026 fact sheet, and EGL principles, producing a reassessment submission "
         "with citations.")
numbered("Plain-language layer. Turns clause text into readable summaries, and works "
         "the other way too — putting a whānau's account into the assessor's terms.")

doc.add_heading("Accessibility", level=2)
para("This is the one idea where accessibility is the product rather than a requirement "
     "to satisfy. WCAG 2.2 AA, screen-reader first, fully keyboard navigable. You have "
     "an accessibility skill installed with reference material.")

doc.add_heading("Build and demo", level=2)
labelled("Why 48 hours works",
         "The simulator is arithmetic over a rules table. The evidence pack is "
         "retrieval over roughly 40 pages of public PDFs. No training, scraping, or "
         "authentication.")
labelled("Demo",
         "Two personas side by side — one whose budget holds, one whose baseline was "
         "suppressed by the 2024 rules. Same tool, different outcomes, both explained. "
         "Then generate the second persona's evidence pack.")

doc.add_heading("Risks", level=2)
bullet("You are modelling contested policy. Label outputs as indicative, not "
       "determinations, and cite every rule to its source.")
bullet("Use synthetic personas only.")
bullet("Get the rules right or do not ship. A wrong number here affects real people, "
       "and saying so in the pitch is worth more than hiding it.")

para("Strongest local resonance, and the only idea where the accessibility brief and "
     "the technical work are the same thing.", italic=True, color=MUTED)

doc.add_page_break()

# ============================== IDEA 4 ==============================
doc.add_heading("Idea 4 — E-waste routing", level=1)

labelled("Problem",
         "New Zealand produces about 99,000 tonnes of e-waste a year under voluntary "
         "schemes only. It is the OECD outlier — producer participation is not "
         "mandatory and there is no extended producer responsibility framework, unlike "
         "every other OECD country. The AS/NZS recycling standards exist but are not "
         "binding.")
labelled("SDG",
         "SDG 12.5 (reduce waste generation). Indicator 12.5.1 is the national "
         "recycling rate, which NZ largely cannot report — and that is part of the "
         "story.")

doc.add_heading("The gap", level=2)
para("E-waste was designated a priority product requiring mandatory stewardship in "
     "2020. TechCollect NZ ran the co-design between 2020 and 2023. As of mid-2026 the "
     "regulations are still not in force and no implementation date is set. All six "
     "2020 priority product groups remain in consultation or drafting. Six years, and "
     "99,000 tonnes a year continues to flow.")

doc.add_heading("Approach", level=2)
numbered("Drop-off guidance. Photograph a device; a vision model identifies the type "
         "and hazardous components, then finds the nearest point that actually accepts "
         "it. Right now this is confusing enough that people give up and landfill it.")
numbered("Repair before recycle. Route repairable devices to repair cafés and "
         "refurbishers first. Recycling is the last step in the waste hierarchy, but "
         "most tools skip straight to it.")
numbered("Reporting ledger. Aggregate anonymised flows into the shape a regulated "
         "scheme will eventually require, so the data model exists when regulations "
         "land.")

doc.add_heading("Build and demo", level=2)
labelled("Low bandwidth",
         "An SMS path that assumes no smartphone makes it usable in places where "
         "informal collectors do this work, which is where 12.5 is weakest globally.")
labelled("Why 48 hours works",
         "Vision identification is one API call. The list of accepting points is small "
         "enough to assemble by hand. The ledger is a schema and a chart.")
labelled("Demo",
         "Photograph an old phone and a dead microwave. One routes to repair, one to a "
         "named recycler, and one is refused by the voluntary scheme entirely — that "
         "refusal gets logged as evidence of the policy gap.")

doc.add_heading("Risks", level=2)
bullet("Frame the regulatory angle as a scheme-ready data model, not a claim to "
       "influence policy.")
bullet("Curate the accepting-points list manually rather than building a scraper.")

para("Best reporting story of the five. Weakest emotional hook, so the pitch has to "
     "carry it.", italic=True, color=MUTED)

doc.add_page_break()

# ============================== IDEA 5 ==============================
doc.add_heading("Idea 5 — Neighbourhood SDG indicator generator", level=1)

labelled("Problem",
         "The main gap in SDG 11 is measurement. Fewer than a third of "
         "sustainable-cities indicators have data at all, and coverage is worst at "
         "neighbourhood level and for vulnerable groups. HLPF 2026 called for urban "
         "data systems as a public good.")
labelled("SDG",
         "SDG 11.7, access to safe and inclusive public space. Indicator 11.7.1 is the "
         "average share of built-up area that is open public space.")

doc.add_heading("The numbers", level=2)
bullet("Access to open public space fell from 48.0% to 45.9% between 2020 and 2025 "
       "across 414 cities in 126 countries. It is going backwards.")
bullet("About 1.1 billion people live in slums or slum-like conditions, with another "
       "2 billion projected over 30 years.")

doc.add_heading("Approach", level=2)
numbered("Compute rather than survey. Free satellite imagery, OpenStreetMap, and a "
         "segmentation model estimate open public space share for a neighbourhood, "
         "output in the form UN-Habitat's methodology expects.")
numbered("Ground-truth layer. Residents contribute phone photos and a vision model "
         "checks whether mapped public space is actually accessible — fenced, locked, "
         "paved over, or usable. Satellites cannot see this.")
numbered("Validate publicly. Run it on a city UN-Habitat has already measured and show "
         "your estimate next to theirs. Judges can check you on the spot, which few "
         "teams offer.")

doc.add_heading("Build and demo", level=2)
labelled("Why 48 hours works",
         "A pre-trained segmentation model, OpenStreetMap via the Overpass API, one "
         "city, one indicator. The discipline is keeping it to exactly that.")
labelled("Demo",
         "Run it on a Hamilton suburb and show the computed figure. Then show the "
         "resident photo layer downgrading a park that turns out to be fenced off, and "
         "explain why the satellite reading was wrong.")

doc.add_heading("Risks", level=2)
bullet("Highest technical risk here. Segmentation quality is unpredictable and hard to "
       "debug late on day two.")
bullet("Pre-compute the demo results and keep a cached fallback.")
bullet("Skip this unless someone on the team has done geospatial or computer vision "
       "work before.")

para("Highest technical ceiling, highest chance of a broken demo.",
     italic=True, color=MUTED)

doc.add_page_break()

# ============================== IDEA 6 ==============================
doc.add_heading("Idea 6 — Patient health literacy and informed consent", level=1)

labelled("Problem",
         "Discharge notes, surgical consent forms, and diagnostic reports are written "
         "in clinical language. Patients sign and leave without understanding what was "
         "decided, what to watch for, or what happens next.")
labelled("SDG",
         "SDG 3 (Good Health), and SDG 10.2 where low literacy tracks with ethnicity "
         "and deprivation. Indicator 3.8.1, coverage of essential health services — "
         "care a patient cannot act on is not effective coverage.")

doc.add_heading("The numbers", level=2)
bullet("More than 56% of New Zealand adults — over 1.6 million people — cannot obtain "
       "and understand basic health information well enough to make informed decisions.")
bullet("Around 90% of Pacific people aged 15 and over have low health literacy. Māori "
       "and Pacific rates are lower than non-Māori, though the largest group by headcount "
       "is Pākehā.")
bullet("16.1% of acute medical admissions among over-65s are readmitted within 30 days, "
       "with higher risk for Māori, Pacific people, and those in deprived areas.")

doc.add_heading("Two corrections to the framing", level=2)
para("Both of these make the pitch stronger, not weaker.", space_after=6)

labelled("It is a legal right, not a service improvement.",
         "Right 5 of the Code of Health and Disability Services Consumers' Rights gives "
         "every consumer the right to effective communication “in a form, language, "
         "and manner that enables the consumer to understand the information provided,” "
         "including a competent interpreter where necessary and reasonably practicable. "
         "Right 6 covers being fully informed. A consent form the patient did not "
         "understand is arguably not compliant. That reframes the tool from helpful to "
         "required.")
labelled("Avoid the word compliance.",
         "New Zealand research consistently frames post-discharge failure as system-side: "
         "inadequate support, medication problems, unpreparedness for community living, "
         "limited whānau support, and poor access to culturally responsive services. "
         "Readmission signals unmet need rather than patient fault. Saying “poor "
         "patient compliance” in front of NZ health judges will cost you.")

doc.add_heading("What already exists", level=2)
para("Epic's patient-facing assistant, Emmie, already answers patient questions inside "
     "MyChart and produces plain-language summaries of lab and procedure results. Do not "
     "pitch a generic medical-jargon translator as though nothing exists.")
para("Two things remain genuinely open. Epic has no shipped feature that rewrites the "
     "full inpatient discharge summary, and New Zealand's public hospitals do not run on "
     "Epic, so none of that reaches a patient at Waikato Hospital. The evidence is also "
     "strong: an LLM-transformed discharge summary moved from an 11th-grade to a "
     "6th-grade reading level, with measured understandability rising from 13% to 81%. "
     "In the same study some errors were introduced, which is why the design below puts "
     "a clinician gate in the middle rather than shipping straight to the patient.")

doc.add_heading("Approach", level=2)
numbered("Ingest and structure. Take a PDF discharge summary or consent form and pull "
         "out the parts a patient actually acts on: medicines and changes, warning signs, "
         "follow-up appointments, activity limits, and who to call.")
numbered("Rewrite with a clinician gate. Generate a plain-language version at a target "
         "reading level, then show it beside the original for a nurse or clinician to "
         "approve or edit before release. The gate is the safety design, and it is worth "
         "saying out loud in the pitch.")
numbered("Visual breakdown. A medication timeline, a red-flag card of symptoms that mean "
         "call someone now, and a follow-up calendar. Most of the comprehension gain "
         "comes from structure, not from prose.")
numbered("Questions on the document. A question box scoped strictly to the patient's own "
         "record, answering only from the approved text and saying so when the answer is "
         "not there. This is the smallest part of the build and should be the smallest "
         "part of the pitch.")
numbered("Language and format options. Te reo Māori, Pacific languages, large print, and "
         "audio. This is where Right 5 is actually met, and it is the part no overseas "
         "product covers for New Zealand.")

doc.add_heading("Positioning", level=2)
para("Health chatbots are the most crowded category at SDG hackathons. If you lead with "
     "the chat box you will look like several other teams. Lead with the consent form and "
     "the legal right, treat the question box as a supporting feature, and the entry reads "
     "differently.")

doc.add_heading("Build and demo", level=2)
labelled("Why 48 hours works",
         "PDF text extraction, one structured extraction call, one rewrite call against a "
         "reading-level target, and a review screen. No training and no integration with "
         "any hospital system. The review interface is ordinary CRUD.")
labelled("Demo",
         "Put a real consent form on screen beside the rewritten version, with the reading "
         "grade shown for each. Then have the clinician view catch and correct one "
         "generated line — demonstrating the gate working is more convincing than "
         "pretending the model is always right.")

doc.add_heading("Risks", level=2)
bullet("Never use real patient records. Use synthetic or published sample documents only, "
       "and say so.")
bullet("The model will occasionally introduce errors. Do not claim otherwise — build the "
       "review step and make it part of the demo.")
bullet("Translation into te reo Māori and Pacific languages needs care. Show it as a "
       "designed pathway with human verification rather than claiming machine translation "
       "is sufficient.")
bullet("Position it as supporting the consent conversation, not replacing it.")

para("Strong legal hook and clear evidence behind it. The risk is presentation — pitched "
     "as a chatbot it blends in, pitched as a Right 5 tool it does not.",
     italic=True, color=MUTED)

doc.add_page_break()

# ============================== COMPARISON ==============================
doc.add_heading("Comparison", level=1)

make_table(
    ["#", "Idea", "SDG", "Novelty", "48h risk", "Emotion", "Tech"],
    [
        ["1", "Cold-chain decisions", "3.b.1", "Medium *", "Low", "High", "Med"],
        ["2", "Grower surplus matching", "12.3 + 2.1", "High", "Low", "High", "Med"],
        ["3", "Disability funding tool", "10.2", "High", "Low", "Highest", "Med"],
        ["4", "E-waste routing", "12.5", "Med–high", "Low", "Low", "Med"],
        ["5", "Indicator generator", "11.7.1", "High", "High", "Medium", "High"],
        ["6", "Health literacy and consent", "3.8.1 + 10.2", "Medium †", "Lowest", "High", "Low"],
    ],
    widths=[0.3, 1.9, 0.95, 0.85, 0.75, 0.8, 0.55],
)
para("* Medium only if pitched as the decision layer. Low if pitched as a logger, "
     "since ColdTrace X already covers that.", size=10, color=MUTED, space_after=2)
para("† Medium because Epic already ships plain-language result summaries. The open "
     "ground is the full discharge summary, the consent form, and te reo Māori and "
     "Pacific language output, none of which overseas products cover for New Zealand.",
     size=10, color=MUTED)

doc.add_heading("Suggested picks", level=2)
make_table(
    ["Strength", "Pick"],
    [
        ["Best all round", "Idea 2 — Grower surplus matching"],
        ["Best local resonance", "Idea 3 — Disability funding tool"],
        ["Best technical showing", "Idea 5 — Indicator generator"],
        ["Safest hardware demo", "Idea 1 — Cold-chain, reframed"],
        ["Easiest to finish well", "Idea 6 — Health literacy and consent"],
    ],
    widths=[2.0, 4.5],
)
para("Idea 6 is the lowest-risk build on the list and the easiest to demo convincingly, "
     "but it sits in the most crowded category. It rewards a team that is stronger on "
     "presentation than on engineering, provided the pitch leads with the Code of Rights "
     "rather than the chat box.")
para("Worth weighing: the probe-in-ice moment in Idea 1 is the most memorable few "
     "seconds available to any team in the room, and a web app cannot match it. If "
     "someone on the team does embedded work, keep the prop — just make sure the "
     "software behind it is the decision layer.")

doc.add_heading("Not included", level=2)
para("Rural volunteer transport matching was considered and dropped. Road XS, Spedsta, "
     "and Assisted Rides already do it, and Hato Hone St John and the Cancer Society "
     "already run the NZ services. Where it has been studied, the bottleneck is human "
     "coordination, not the matching software.")

doc.add_heading("Practical notes", level=2)
para("Teams split at noon on day two — half keep building, half prepare the pitch. "
     "Agree who is on which side beforehand. The 30 minutes between the end of the hack "
     "and the start of pitching is the only window to regroup.")
para("Pitch structure, from the official template:", space_after=4)
make_table(
    ["Time", "Section", "Content"],
    [
        ["0:30", "Connect", "Who you are, what you care about"],
        ["1:00", "Problem", "Who is affected, made human, with data"],
        ["2:00", "Big idea", "How it solves the problem, plus the prototype"],
        ["1:30", "Impact", "What changes, what you need, what it means longer term"],
    ],
    widths=[0.6, 1.1, 4.8],
)
para("The template notes: describe the problem you are solving, which may not be the "
     "one you set out to solve.", italic=True, color=MUTED)

doc.add_page_break()

# ============================== SOURCES ==============================
doc.add_heading("Sources", level=1)

SOURCES = [
    ("SDG", [
        ("UN SDG Report 2026", "https://unstats.un.org/sdgs/report/2026/"),
        ("SDSN Report 2026", "https://dashboards.sdgindex.org/"),
        ("UN-Habitat SDG 11 Report 2026",
         "https://unhabitat.org/sdg-11-global-report-2026-housing-at-the-centre-of-sustainable-cities-and-communities"),
    ]),
    ("Cold chain", [
        ("Health NZ cold chain standards",
         "https://www.tewhatuora.govt.nz/health-services-and-programmes/vaccine-information/vaccine-service-delivery/cold-chain-standards-for-vaccines/"),
        ("National Standards 2017 (2nd ed)",
         "https://www.tewhatuora.govt.nz/assets/For-the-health-sector/Health-sector-guidance/Vaccine-information-for-healthcare-professionals/National-Standards-for-Vaccine-Storage-and-Transportation-for-Immunisation-Providers-2017-2nd-edition.pdf"),
        ("IMAC cold chain", "https://immune.org.nz/vaccines/cold-chain"),
        ("Nexleaf ColdTrace", "https://www.nexleaf.org/work/vaccines/"),
    ]),
    ("Food", [
        ("NZFN Hunger Monitor 2025",
         "https://www.nzfoodnetwork.org.nz/wp-content/uploads/2026/03/NZFN-Hunger-Monitor-2025-Digital.pdf"),
        ("NZFN impact stats", "https://www.nzfoodnetwork.org.nz/our-impact/"),
        ("NZFN and Lineage partnership",
         "https://www.nzfoodnetwork.org.nz/food-rescue-partnership-nz-food-network-lineage-grant/"),
        ("Rabobank food waste", "https://www.rabobank.co.nz/community/food-waste"),
        ("PMCSA food rescue evidence", "https://www.pmcsa.ac.nz/topics/food-rescue-food-waste/"),
    ]),
    ("Disability", [
        ("DSS February 2026 fact sheet",
         "https://www.disabilitysupport.govt.nz/disabled-people/improvements-to-disability-support-services/february-2026-fact-sheet"),
        ("NZ Disability Strategy 2026–2030",
         "https://www.whaikaha.govt.nz/assets/NZDS-26-30-Documents/NZDS-2026-2030.pdf"),
        ("CCS Disability Action, Budget 2026",
         "https://www.ccsdisabilityaction.org.nz/news/budget-2026-statement"),
        ("Blind Low Vision NZ submission",
         "https://blindlowvision.org.nz/news/submission-on-the-disability-support-services-bill-2026/"),
    ]),
    ("E-waste", [
        ("TechCollect NZ stewardship",
         "https://techcollect.nz/e-waste-product-stewardship-new-zealand/"),
        ("Electronic waste in New Zealand",
         "https://en.wikipedia.org/wiki/Electronic_waste_in_New_Zealand"),
    ]),
    ("Data and tooling", [
        ("data.govt.nz APIs", "https://data.govt.nz/catalogue-guide/using-data-govt-nz-apis"),
        ("NZ AI Strategy",
         "https://www.mbie.govt.nz/assets/new-zealands-strategy-for-artificial-intelligence.pdf"),
    ]),
    ("Health literacy and consent", [
        ("HDC Code of Rights (Rights 5 and 6)",
         "https://www.hdc.org.nz/your-rights/about-the-code/code-of-health-and-disability-services-consumers-rights/"),
        ("Code of Rights, full PDF",
         "https://www.hdc.org.nz/media/550hs5ih/code-of-rights_online_5-sept-2022.pdf"),
        ("HQSC health literacy background",
         "https://www.hqsc.govt.nz/assets/Consumer-hub/Partners-in-Care/Publications-resources/health-literacy-background-info-Sep-2013.pdf"),
        ("Health literacy and Pacific peoples review",
         "https://ojs.aut.ac.nz/pacific-health/article/view/4"),
        ("NZ Society of Anaesthetists on health literacy",
         "https://anaesthesia.nz/news/health-literacy-in-aotearoa-new-zealand-remains-a-critical-challenge/"),
        ("AI-generated patient-friendly discharge summaries",
         "https://www.medrxiv.org/content/10.1101/2025.07.04.25330804.full.pdf"),
        ("NYU Langone discharge summary trial",
         "https://clinicaltrials.gov/study/NCT06711458"),
        ("Epic Emmie, patient-facing AI", "https://www.epic.com/software/emmie/"),
        ("Acute readmissions to hospitals, MoH",
         "https://www.health.govt.nz/statistics-research/system-monitoring/planning-and-performance-data/acute-readmissions-to-hospitals"),
    ]),
    ("Judging", [
        ("GitLab AI Hackathon 2026 winners",
         "https://about.gitlab.com/blog/gitlab-ai-hackathon-2026-meet-the-winners/"),
    ]),
]

for group, items in SOURCES:
    doc.add_heading(group, level=3)
    for label, url in items:
        source_line(label, url)

doc.add_heading("Technical notes", level=2)
para("data.govt.nz documentation references CKAN 2.7, which is old. Check "
     "/api/3/action/status_show before coding against those docs. The DataStore API "
     "only exists where a clean CSV was supplied; many listings are metadata-only links "
     "to agency sites.")
para("Claude Sonnet 5 pricing rises to $3/$15 per million tokens after 31 August 2026.")

out = r"C:\Users\ericc\OneDrive\Desktop\Programs\Hackathon-Project\AI Hackathon 2026 - Ideation.docx"
doc.save(out)
print("saved:", out)
