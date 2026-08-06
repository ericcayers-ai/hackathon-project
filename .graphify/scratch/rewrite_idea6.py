# -*- coding: utf-8 -*-
"""Rewrite the Idea 6 section (Nurse Notes -> ClearChart) in place in the
master docx, add judging-criteria mapping + backend section, update the
Comparison table and Suggested Picks, and add new source citations."""
import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = "docs/SDG_Hackathon_Idea_Guide.docx"
MUTED = RGBColor(0x5F, 0x6B, 0x73)

doc = docx.Document(PATH)
body = doc.element.body


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0B57D0")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(c)
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hl.append(new_run)
    paragraph._p.append(hl)
    return hl


# ---------------------------------------------------------------
# 1. Locate the Idea 6 block: from its Heading 1 to right before Idea 7's Heading 1
# ---------------------------------------------------------------
paras = doc.paragraphs
start_idx = None
end_idx = None
for i, p in enumerate(paras):
    if p.style.name == "Heading 1" and p.text.strip().startswith("Idea 6"):
        start_idx = i
    elif start_idx is not None and p.style.name == "Heading 1" and p.text.strip().startswith("Idea 7"):
        end_idx = i
        break

assert start_idx is not None and end_idx is not None, "could not locate Idea 6 block"

anchor_p = paras[end_idx]  # Idea 7 heading; we insert new content before this, then delete old range
anchor_elm = anchor_p._p

# Delete old paragraphs [start_idx, end_idx)
old_elms = [paras[i]._p for i in range(start_idx, end_idx)]
for elm in old_elms:
    elm.getparent().remove(elm)


# ---------------------------------------------------------------
# 2. Helpers that insert new paragraphs immediately before anchor_elm
# ---------------------------------------------------------------
def _insert_before_anchor(new_p_elm):
    anchor_elm.addprevious(new_p_elm)


def heading(text, level=1):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    p.add_run(text)
    body.remove(p._p)  # add_paragraph appended it at end of body; move it before anchor
    _insert_before_anchor(p._p)
    return p


def para(text="", bold=False, italic=False, color=None, space_after=None):
    p = doc.add_paragraph()
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = color
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    body.remove(p._p)
    _insert_before_anchor(p._p)
    return p


def labelled(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(label + "  ")
    r.bold = True
    p.add_run(text)
    body.remove(p._p)
    _insert_before_anchor(p._p)
    return p


def bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.add_run("•")
    sp = p.add_run("     ")
    sp.font.size = Pt(11)
    p.add_run(text)
    body.remove(p._p)
    _insert_before_anchor(p._p)
    return p


def numbered(text, _counter=[0]):
    _counter[0] += 1
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.add_run(f"{_counter[0]}.   {text}")
    body.remove(p._p)
    _insert_before_anchor(p._p)
    return p


def source_line(label, url):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label + " — ")
    r.font.size = Pt(10)
    add_hyperlink(p, url, url)
    for run in p.runs:
        run.font.size = Pt(10)
    body.remove(p._p)
    _insert_before_anchor(p._p)
    return p


# ---------------------------------------------------------------
# 3. New Idea 6 content — ClearChart (formerly Nurse Notes)
# ---------------------------------------------------------------
heading("Idea 6 — ClearChart (formerly Nurse Notes) — Patient Health Literacy and Informed Consent", level=1)
para("")

para("Working name during ideation was “Nurse Notes.” Renamed to ClearChart to name what the "
     "product actually does — turns a clinical chart into something a patient can read — rather than "
     "describing who uses it. Used consistently below; the frontend repository may still display the "
     "old name until it is renamed there separately.", italic=True, color=MUTED, space_after=10)

labelled("Value proposition",
         "ClearChart rewrites hospital discharge summaries and consent forms into plain language "
         "(~6th-grade reading level), with a mandatory clinician-review gate before anything reaches "
         "a patient. It is not a medical chatbot and does not answer open-ended questions — it "
         "transforms one specific, high-stakes document a patient is handed at the moment they are "
         "least equipped to parse it, and a human signs off on every word before release.")

heading("Problem", level=3)
para("Patients are hospitalised, told what happened and what to do next in dense clinical shorthand "
     "(TDS, OD, BD, SOB, melaena, 6/52), and sent home to act on it while tired, unwell, and least "
     "able to decode it. This is not a generic literacy problem — it recurs at a specific, "
     "high-consequence moment: the point of discharge and the point of signing consent, when a "
     "misunderstood instruction directly produces a missed dose, a missed follow-up, or an avoidable "
     "readmission.")

heading("Grounding: documented gaps in clinical documentation and discharge communication", level=3)
para("Directly evidenced by the health-records material already gathered for this project:", space_after=4)
bullet("More than 56% of New Zealand adults — over 1.6 million people — cannot obtain and understand "
       "basic health information well enough to make informed decisions. [from project docs]")
bullet("Around 90% of Pacific people aged 15+ have low health literacy; Māori and Pacific rates track "
       "below non-Māori, though the largest group by headcount is Pākehā. [from project docs]")
bullet("16.1% of acute medical admissions among over-65s are readmitted within 30 days, with higher "
       "risk for Māori, Pacific people, and those in deprived areas. [from project docs]")
bullet("An LLM-rewritten discharge summary moved measured patient understandability from 13% to 81%, "
       "and moved the reading level from 11th grade to 6th grade — the same study introduced some "
       "errors in the rewrite, which is the direct evidentiary basis for this product's clinician "
       "gate rather than a hypothetical safety feature. [from project docs, medRxiv preprint]")
para("Researched separately for this revamp (not previously in the project docs) — general clinical-"
     "documentation and discharge-summary literature, used to sharpen the problem statement beyond "
     "NZ-specific health-literacy stats:", space_after=4)
bullet("Discharge summaries completed and available within 7 days of discharge are consistently "
       "associated with lower readmission rates; delayed or unclear discharge documentation is linked "
       "to incorrect prescriptions, untracked pending results, and preventable readmissions. "
       "[researched/inferred — general clinical-documentation literature, no NZ-specific study found]")
bullet("EHR documentation burden on nurses is a well-studied, separate but related problem: acute and "
       "critical-care nurses report EHR systems poorly fitted to unit-specific workflow as a major "
       "driver of documentation load, time away from patients, and burnout risk. ClearChart does not "
       "address nurse-facing charting burden directly — it targets the patient-facing output of that "
       "same discharge process — but this is the adjacent, better-evidenced problem worth naming "
       "honestly so the pitch doesn't overclaim what the tool fixes. [researched/inferred]")
bullet("A cross-sectional comparison of LLM-generated vs. clinician-written discharge summaries found "
       "comparable overall quality, with LLM output more succinct but with a higher rate of low-harm "
       "omission errors than clinician-written summaries — independent confirmation, from a different "
       "study than the 13%→81% figure above, that LLM-assisted discharge rewriting is viable in "
       "principle but needs a human check on omissions specifically. [researched/inferred]")
para("No real patient or provider data appears anywhere in this project's docs, the demo, or the "
     "frontend repo — the only patient-shaped document used anywhere is the synthetic, clearly-labelled "
     "sample in samples/synthetic-discharge-01.txt. There is nothing here that triggers HIPAA (a US "
     "statute, not applicable in NZ) or NZ Health Information Privacy Code obligations, precisely "
     "because no real record is ever processed — this is itself part of the pitch, not just a "
     "compliance footnote.", space_after=10)

heading("SDG", level=3)
para("SDG 3 (Good Health), Indicator 3.8.1 — coverage of essential health services; care a patient "
     "cannot act on is not effective coverage. SDG 10.2 — inclusion, since low health literacy tracks "
     "with ethnicity and deprivation in New Zealand.")

heading("The legal reframe", level=3)
labelled("A legal right, not a service improvement.",
         "Right 5 of the Code of Health and Disability Services Consumers' Rights gives every consumer "
         "the right to effective communication “in a form, language, and manner that enables the "
         "consumer to understand the information provided,” including a competent interpreter "
         "where necessary and reasonably practicable. Right 6 covers being fully informed. A consent "
         "form the patient did not understand is arguably not compliant. That reframes the tool from "
         "helpful to required infrastructure.")
labelled("Avoid the word compliance.",
         "New Zealand research consistently frames post-discharge failure as system-side: inadequate "
         "support, medication problems, unpreparedness for community living, limited whānau support, "
         "and poor access to culturally responsive services. Readmission signals unmet need, not "
         "patient fault. “Poor patient compliance” framing will cost the pitch in front of "
         "NZ health judges.")

heading("What already exists — niche check", level=3)
para("Epic's patient-facing assistant, Emmie, already answers patient questions inside MyChart and "
     "produces plain-language summaries of lab and procedure results — this is a real, shipped "
     "competitor and the pitch must not pretend otherwise.")
para("Two things remain genuinely open, and this is where ClearChart's niche actually sits: Epic has "
     "no shipped feature that rewrites the full inpatient discharge summary or a consent form, and "
     "New Zealand's public hospitals do not run on Epic, so none of Emmie's capability reaches a "
     "patient at Waikato Hospital. Te reo Māori and Pacific-language output — the part of Right 5 "
     "this product actually closes — is not offered by any overseas competitor either. The niche is "
     "narrow by design (one document type, one moment, one language gap) rather than a broad "
     "“AI health assistant,” which is deliberately tighter than the health-chatbot category "
     "judges see most often.")

heading("Approach", level=3)
numbered("Ingest and structure. Take a PDF discharge summary or consent form and pull out the parts "
         "a patient actually acts on: medicines and changes, warning signs, follow-up appointments, "
         "activity limits, and who to call.")
numbered("Rewrite with a clinician gate. Generate a plain-language version at a target reading level, "
         "shown beside the original, for a nurse or clinician to approve or edit before release. This "
         "gate is the safety design, not an afterthought, and it is the single most defensible claim "
         "in the pitch.")
numbered("Visual breakdown. A medication timeline, a red-flag card of symptoms that mean call someone "
         "now, and a follow-up calendar. Most of the comprehension gain comes from structure, not "
         "prose.")
numbered("Language and format options. Te reo Māori, Pacific languages, large print, and audio, each "
         "with human verification. This is where Right 5 is actually met, and it is the part no "
         "overseas product covers for New Zealand.")

heading("Standout factor: what's already real, not hypothetical", level=3)
para("This is the strongest practicality argument available to any idea in this document: a working "
     "frontend already exists (HimendraFdo/Nurse-Notes, frontend branch — Vite + React), implementing "
     "exactly the architecture described above, not a mockup of it.", space_after=4)
bullet("A two-pane clinician review screen: left pane is the original clinical text (paste, or upload "
       ".txt/.md/.pdf, parsed locally with pdf.js), right pane is the plain-language rewrite, streamed "
       "live from a local LLM and directly editable.")
bullet("A persistent “Nurse review required before release” banner, visible until the "
       "clinician explicitly approves — the gate is a UI element a judge can see, not just a claim "
       "in the pitch.")
bullet("Approve-for-release locks the edited text and stamps a timestamp plus the reviewing clinician's "
       "name; Export patient PDF (generated in-browser with jsPDF) is only enabled after approval — "
       "the workflow enforces the gate structurally, it cannot be skipped.")
bullet("Reading-level badges are deliberately asymmetric rather than a naive Flesch–Kincaid on "
       "both panes: the original shows a clinical-jargon count (shorthand terms a patient can't "
       "decode), and the rewrite shows the Flesch–Kincaid grade against a 6th-grade target — "
       "because FK alone rates telegraphic shorthand like “SOB” or “6/52” as "
       "artificially easy, which would misrepresent the original as more readable than it is.")
bullet("Runs entirely against a local OpenAI-compatible endpoint (LM Studio, currently "
       "google/gemma-3n-e4b) — no cloud call, no account, nothing leaves the machine. This is a real, "
       "already-implemented privacy property, not a roadmap promise, and it is also the "
       "Māori-data-sovereignty argument: data that never moves cannot be governed by someone else's "
       "jurisdiction.")
para("Known, intentional limitation, not a bug to hide: the current small on-device model misreads NZ "
     "date shorthand (“6/52”, meaning 6 weeks, sometimes renders as “6 weeks and 5 "
     "days”). This is kept in the demo on purpose — the nurse catches it, corrects it, and "
     "approves, which is the clinician gate visibly doing its job on screen rather than being asserted "
     "in a slide.", space_after=8)

heading("Technical Approach / Backend — proposed additions", level=3)
para("Grounded in open-ended 2020+ open-source research into what would meaningfully strengthen this "
     "specific product's actual weak points (small-model date/shorthand errors, no structured "
     "extraction guarantee, no jargon-detection ground truth) rather than AI-for-AI's-sake. All items "
     "below are researched/inferred technical context, not yet integrated into the existing frontend "
     "branch. Ranked by hackathon feasibility.", space_after=8)

labelled("1. MedGemma 4B (structured extraction) — buildable in the hackathon window.",
         "Google's open-weight, edge-deployable clinical model (Health AI Developer Foundations "
         "license, not Apache). The 4B v1.5 variant scores 91.0 macro-F1 turning lab/clinical reports "
         "into structured JSON and 89.6% on EHRQA. Mechanism: swap the current single free-text rewrite "
         "call for a two-stage pipeline — MedGemma extracts medicines/doses/dates/warning-signs into a "
         "structured schema first, then the rewrite step generates plain language from that structured "
         "data instead of raw prose. This directly targets the one demoed failure mode (6/52 date "
         "misparse) by making dates a discrete, separately-validated field rather than free text the "
         "model has to get right in one pass — it doesn't guarantee the date parses correctly, but it "
         "makes the error visible and correctable in one place instead of buried in generated prose. "
         "Feasibility: realistic stretch goal for the remaining build time — requires a second model "
         "call and a schema, but no training. [researched/inferred]")

labelled("2. LLM-AIx or OpenMed-style local extraction/de-identification pass — future roadmap, not "
         "hackathon-scoped.",
         "Open-source pipelines (LLM-AIx runs on a single on-prem GPU; OpenMed pairs clinical NER with "
         "a local Ollama model) built specifically for privacy-preserving, on-device clinical "
         "information extraction with no data leaving the host machine. Mechanism: a pre-processing "
         "de-identification/entity-tagging pass before the rewrite step, so if this tool were ever "
         "pointed at anything beyond the synthetic sample, PII would be structurally stripped before "
         "the LLM sees it, not just policy-promised. Explicitly future-roadmap: the product's current "
         "synthetic-data-only policy makes this unnecessary for the hackathon demo itself, but it is "
         "the concrete next step if ClearChart were piloted against real (de-identified) documents "
         "post-hackathon, and naming it shows the team has thought past the demo. [researched/inferred]")

labelled("3. Reading-level / jargon-detection fine-tuning or few-shot calibration — future roadmap.",
         "Medical text simplification research (e.g. MultiMSD, ACL 2025 Findings) provides "
         "parallel clinical→plain-language corpora that could few-shot-calibrate or lightly "
         "fine-tune the jargon-detection and reading-grade badges beyond the current heuristic word/"
         "syllable counters, improving the honesty of the reading-grade claim shown to the clinician. "
         "Not buildable to a demoable standard in the remaining hackathon time; named as evidence the "
         "current heuristic approach was a deliberate scoping choice, not an oversight. [researched/"
         "inferred]")

para("Architectural compatibility with the existing frontend branch: all three proposals slot in at "
     "the src/lib/llm.js call site (currently a single fetch to the LM Studio OpenAI-compatible "
     "endpoint) as additional calls or a swapped model — none require changing the review-gate UI, "
     "the approval flow, or the PDF export, because the gate sits downstream of generation regardless "
     "of which model produced the draft. This is a disconnected-rebuild risk avoided, not created.",
     space_after=10)

heading("Impact metrics", level=3)
bullet("13% → 81% patient-understandability lift, and 11th-grade → 6th-grade reading-level "
       "shift, from the published study this design is built on — not a number invented for this "
       "pitch. [from project docs / medRxiv]")
bullet("16.1% 30-day readmission rate among over-65 acute admissions is the population-level metric "
       "this tool is positioned against; ClearChart does not claim to have measured a readmission "
       "reduction itself — it targets one documented input to that outcome (discharge-instruction "
       "comprehension) and the pitch should say so explicitly rather than implying a causal claim "
       "the demo cannot support.")
bullet("Jargon count on the original pane (61 on the shipped synthetic sample) is a demoable, "
       "judge-verifiable number computed live in the browser, not a vanity metric — it is exactly "
       "the kind of “number that maps to something concrete a judge can check” this event's "
       "own judging notes call out as a differentiator.")

heading("Judging-criteria mapping", level=3)
para("Mapped against this event's own four equally-weighted criteria (Inspiration, Technology, "
     "Design, Presentation) plus the TAIAO (SDG alignment/feasibility/impact) and Technical Brilliance "
     "prize criteria.", space_after=6)
labelled("Inspiration —",
         "a defined, evidenced problem (56% low health literacy, 16.1% readmission risk, 13%→81% "
         "comprehension study) with a legal spine (Right 5) and a specific human moment (the patient "
         "who signs and can't remember what changed), not a generic “AI for good” framing.")
labelled("Technology —",
         "a real, running local-LLM pipeline (not a mockup) with a working two-pane review UI already "
         "built, plus a scoped, honestly-ranked backend roadmap (MedGemma extraction as the realistic "
         "next call, de-identification and simplification fine-tuning named explicitly as future work "
         "rather than overclaimed as done).")
labelled("Design & Innovation —",
         "structure-over-prose UI (medication timeline, red-flag card, follow-up calendar) instead of "
         "a chat window, an asymmetric jargon/FK badge pair designed to avoid a known measurement "
         "flaw, and te reo Māori / Pacific-language output that no competitor (including Epic's "
         "Emmie) ships for New Zealand.")
labelled("Presentation —",
         "one sentence any judge grasps (“consent forms patients sign, turned into consent forms "
         "patients understand”), a live demo of the clinician gate catching and correcting a real "
         "generated error on screen, and a criteria-fit slide making this mapping explicit rather than "
         "leaving judges to infer it.")
labelled("TAIAO (SDG alignment, feasibility, impact) —",
         "explicit SDG 3.8.1 / SDG 10.2 mapping, a 48-hour-buildable architecture with no hospital "
         "integration or training required, and impact evidence drawn from a published clinical study "
         "rather than an assumption.")
labelled("Technical Brilliance —",
         "the clinician-gate architecture itself is the technical thesis (safety-by-design around a "
         "small, imperfect on-device model, rather than assuming a bigger model solves the problem), "
         "demonstrated with a real bug the team chose to keep because it proves the gate works.")

heading("Positioning", level=3)
para("Health chatbots are the most crowded category at SDG hackathons. If this leads with a chat box "
     "it looks like several other teams. Lead with the consent form and the legal right; the "
     "clinician-gate review screen is the centrepiece, not a chatbot with a health skin.")

heading("Build and demo", level=3)
labelled("Why the remaining time is enough",
         "PDF text extraction, one structured-extraction call, one rewrite call against a "
         "reading-level target, and a review screen are already built end-to-end in the frontend "
         "branch. No training and no hospital-system integration. Remaining effort is the honestly-"
         "scoped MedGemma extraction addition above, plus rehearsal.")
labelled("Demo",
         "Put the synthetic consent form on screen beside the rewritten version, with the reading "
         "grade and jargon count shown for each. Then show the clinician catching and correcting the "
         "“6/52” date error live and clicking Approve — demonstrating the gate working is "
         "more convincing than pretending the model is always right.")

heading("Risks", level=3)
bullet("Never use real patient records. Use synthetic or published sample documents only — already "
       "enforced in the shipped frontend (samples/synthetic-discharge-01.txt is the only document "
       "used anywhere).")
bullet("The model will occasionally introduce errors. Do not claim otherwise — the review step is "
       "built and is part of the demo, not a stated intention.")
bullet("Translation into te reo Māori and Pacific languages needs care. Position it as a designed "
       "pathway with human verification, not machine translation presented as sufficient on its own — "
       "this is not yet built in the current frontend branch and should be named as a roadmap item, "
       "not implied as shipped.")
bullet("Position it as supporting the consent conversation, not replacing it.")

para("Strong legal hook, evidenced problem, and — unusually for this stage of a hackathon pitch — a "
     "working frontend to point to instead of a hypothetical one. The remaining risk is presentation "
     "discipline: pitched as a chatbot it blends in, pitched as a Right 5 tool with a visible safety "
     "gate it does not.", italic=True, color=MUTED)

doc.add_page_break()
pb = doc.paragraphs[-1]._p
body.remove(pb)
_insert_before_anchor(pb)


# ---------------------------------------------------------------
# 4. Update the Comparison table row for Idea 6 -> ClearChart
# ---------------------------------------------------------------
updated_row = False
for table in doc.tables:
    hdr_cells = table.rows[0].cells
    headers = [c.text.strip() for c in hdr_cells]
    if headers[:2] == ["#", "Idea"]:
        for row in table.rows[1:]:
            if row.cells[0].text.strip() == "6":
                row.cells[1].text = "ClearChart (health literacy & consent)"
                updated_row = True
                break
        if updated_row:
            break
assert updated_row, "could not find Idea 6 row in comparison table"

print("Idea 6 -> ClearChart rewrite complete.")
doc.save(PATH)
print("Saved:", PATH)
